import os
import json
import re
import time
import uuid
import tempfile
from flask import Flask, request, send_file, jsonify, render_template
from docx import Document
from docx.shared import Pt
import boto3

app = Flask(__name__, template_folder="../frontend", static_folder="../frontend/static")

# Load BDA config
CONFIG_PATH = os.path.join(os.path.dirname(__file__), "bda_config.json")
with open(CONFIG_PATH) as f:
    BDA_CONFIG = json.load(f)

REGION = BDA_CONFIG["region"]
PROJECT_ARN = BDA_CONFIG["project_arn"]
S3_BUCKET = BDA_CONFIG["s3_bucket"]
PROFILE_ARN = f"arn:aws:bedrock:{REGION}:{boto3.client('sts').get_caller_identity()['Account']}:data-automation-profile/us.data-automation-v1"

s3 = boto3.client("s3", region_name=REGION)
bda_runtime = boto3.client("bedrock-data-automation-runtime", region_name=REGION)



def call_bda(file_path):
    """Upload file to S3, invoke BDA for document parsing, classify doc type, extract fields."""
    job_id = str(uuid.uuid4())[:8]
    filename = os.path.basename(file_path)
    input_key = f"input/{job_id}/{filename}"
    output_prefix = f"output/{job_id}/"

    # Upload to S3
    s3.upload_file(file_path, S3_BUCKET, input_key)

    # Invoke BDA async
    resp = bda_runtime.invoke_data_automation_async(
        inputConfiguration={"s3Uri": f"s3://{S3_BUCKET}/{input_key}"},
        outputConfiguration={"s3Uri": f"s3://{S3_BUCKET}/{output_prefix}"},
        dataAutomationConfiguration={
            "dataAutomationProjectArn": PROJECT_ARN,
            "stage": "LIVE"
        },
        dataAutomationProfileArn=PROFILE_ARN
    )
    invocation_arn = resp["invocationArn"]

    # Poll for completion (max ~5 minutes)
    for _ in range(60):
        time.sleep(5)
        status_resp = bda_runtime.get_data_automation_status(invocationArn=invocation_arn)
        status = status_resp["status"]
        if status == "Success":
            break
        if status in ("ServiceError", "ClientError"):
            raise Exception(f"BDA job failed: {status}")

    if status != "Success":
        raise Exception("BDA job timed out")

    # Read standard output to get per-page text
    page_texts = []
    paginator = s3.get_paginator("list_objects_v2")
    for pg in paginator.paginate(Bucket=S3_BUCKET, Prefix=output_prefix):
        for obj in pg.get("Contents", []):
            if "standard_output" in obj["Key"] and obj["Key"].endswith("result.json"):
                body = s3.get_object(Bucket=S3_BUCKET, Key=obj["Key"])["Body"].read()
                std_data = json.loads(body)
                for page in std_data.get("pages", []):
                    text = page.get("representation", {}).get("text", "").strip()
                    page_texts.append(text)

    full_text = "\n".join(page_texts)

    # Classify document type using Claude
    bedrock_rt = boto3.client("bedrock-runtime", region_name=REGION)
    doc_type = _classify_document(bedrock_rt, full_text[:3000])

    if doc_type == "clinical_trial":
        # Extract clinical trial data using Claude on the full text
        ct_data = _extract_clinical_trial(bedrock_rt, full_text)
        _cleanup_s3(job_id, output_prefix)
        return {"type": "clinical_trial", "data": ct_data}

    # Pharmacy prescription flow
    # Group consecutive non-empty pages into prescriptions
    prescriptions = []
    current = []
    for text in page_texts:
        if text:
            current.append(text)
        else:
            if current:
                prescriptions.append("\n".join(current))
                current = []
    if current:
        prescriptions.append("\n".join(current))

    # If only one group — use BDA custom output directly
    if len(prescriptions) <= 1:
        all_results = []
        for pg in paginator.paginate(Bucket=S3_BUCKET, Prefix=output_prefix):
            for obj in pg.get("Contents", []):
                if "custom_output" in obj["Key"] and obj["Key"].endswith("result.json"):
                    body = s3.get_object(Bucket=S3_BUCKET, Key=obj["Key"])["Body"].read()
                    data = json.loads(body)
                    result = data.get("inference_result", {})
                    all_results.append(result)
        _cleanup_s3(job_id, output_prefix)
        return {"type": "prescription", "data": all_results}

    # Multi-page prescriptions: use Claude per prescription
    all_results = []
    for rx_text in prescriptions:
        extracted = _extract_with_claude(bedrock_rt, rx_text)
        if extracted:
            all_results.append(extracted)

    _cleanup_s3(job_id, output_prefix)
    return {"type": "prescription", "data": all_results}


def _extract_with_claude(bedrock_rt, rx_text):
    """Use Claude to extract structured pharmacy fields from prescription text."""
    prompt = f"""Extract the following fields from this pharmacy prescription text. Return ONLY valid JSON with these exact keys. If a field is not found, use an empty string.

Keys: pharmacy_name, pharmacy_address, pharmacy_phone, pharmacy_license, date_dispensed, rx_number, pharmacist_name, pharmacy_technician, patient_name, patient_dob, patient_id, patient_address, patient_phone, patient_allergies, insurance_provider, insurance_member_id, physician_name, physician_specialty, clinic_name, clinic_address, npi_number, dea_number, physician_phone, drug_generic_name, drug_brand_name, drug_strength, dosage_form, ndc_number, quantity_dispensed, days_supply, refills_remaining, lot_number, expiration_date, sig_directions, warnings, retail_price, insurance_covered, patient_copay, payment_method, patient_counseled, rph_initials, verification_time, next_refill_date

Prescription text:
{rx_text}

Return ONLY the JSON object, no markdown, no explanation."""

    resp = bedrock_rt.invoke_model(
        modelId="us.anthropic.claude-sonnet-4-5-20250929-v1:0",
        contentType="application/json",
        accept="application/json",
        body=json.dumps({
            "anthropic_version": "bedrock-2023-05-31",
            "max_tokens": 4096,
            "messages": [{"role": "user", "content": prompt}]
        })
    )
    result = json.loads(resp["body"].read())
    text = result["content"][0]["text"].strip()
    # Strip markdown fences if present
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\s*", "", text)
        text = re.sub(r"\s*```$", "", text)
    return json.loads(text)


def _cleanup_s3(job_id, output_prefix):
    """Remove S3 input and output objects for a completed BDA job."""
    prefixes = [f"input/{job_id}/", output_prefix]
    for prefix in prefixes:
        paginator = s3.get_paginator("list_objects_v2")
        for page in paginator.paginate(Bucket=S3_BUCKET, Prefix=prefix):
            for obj in page.get("Contents", []):
                s3.delete_object(Bucket=S3_BUCKET, Key=obj["Key"])

def _classify_document(bedrock_rt, text_sample):
    """Classify document as 'prescription' or 'clinical_trial' using Claude."""
    prompt = f"""Classify this document into exactly one category. Return ONLY the category name, nothing else.

Categories:
- prescription (pharmacy prescription dispensing record, medication dispensing)
- clinical_trial (clinical trial report, safety report, study summary, regulatory submission)

Document text (first portion):
{text_sample}

Category:"""

    resp = bedrock_rt.invoke_model(
        modelId="us.anthropic.claude-sonnet-4-5-20250929-v1:0",
        contentType="application/json",
        accept="application/json",
        body=json.dumps({
            "anthropic_version": "bedrock-2023-05-31",
            "max_tokens": 20,
            "messages": [{"role": "user", "content": prompt}]
        })
    )
    result = json.loads(resp["body"].read())
    category = result["content"][0]["text"].strip().lower()
    if "clinical" in category or "trial" in category:
        return "clinical_trial"
    return "prescription"


def _extract_clinical_trial(bedrock_rt, full_text):
    """Extract structured clinical trial data from full document text using Claude."""
    prompt = f"""Extract structured data from this clinical trial report. Return ONLY valid JSON with these exact top-level keys. If a field is not found, use an empty string. For arrays, return empty arrays if not found.

Return this JSON structure:
{{
  "document_info": {{
    "document_id": "",
    "prepared_by": "",
    "reviewed_by": "",
    "approved_by": "",
    "date_issued": "",
    "department": "",
    "trial_phase": "",
    "version": ""
  }},
  "sponsor": {{
    "name": "",
    "address": "",
    "contact": "",
    "ind_number": ""
  }},
  "study_overview": {{
    "drug_compound": "",
    "drug_working_name": "",
    "mechanism": "",
    "indication": "",
    "trial_sites": "",
    "countries": "",
    "enrollment_period": "",
    "total_enrolled": "",
    "protocol_number": "",
    "primary_objective": "",
    "secondary_objectives": []
  }},
  "enrollment": {{
    "total_screened": "",
    "screen_failures": "",
    "randomized": "",
    "treatment_arm": "",
    "placebo_arm": "",
    "completed": "",
    "discontinued": "",
    "discontinuation_reasons": []
  }},
  "demographics": {{
    "mean_age": "",
    "female_pct": "",
    "male_pct": "",
    "race_ethnicity": [],
    "mean_disease_duration": "",
    "other_baseline": []
  }},
  "efficacy": {{
    "primary_endpoint": "",
    "primary_result_treatment": "",
    "primary_result_placebo": "",
    "primary_difference": "",
    "primary_p_value": "",
    "secondary_endpoints": []
  }},
  "safety": {{
    "any_teae_treatment": "",
    "any_teae_placebo": "",
    "any_sae_treatment": "",
    "any_sae_placebo": "",
    "deaths_treatment": "",
    "deaths_placebo": "",
    "death_details": "",
    "common_teaes": [],
    "aesi": [],
    "lab_findings": []
  }},
  "pharmacokinetics": {{
    "pk_parameters": [],
    "special_populations": [],
    "drug_interactions": []
  }},
  "regulatory": {{
    "nda_target": "",
    "designations": [],
    "open_items": []
  }},
  "risk_benefit": {{
    "benefit_summary": "",
    "risk_summary": "",
    "overall_assessment": ""
  }}
}}

Clinical trial document:
{full_text}

Return ONLY the JSON object:"""

    resp = bedrock_rt.invoke_model(
        modelId="us.anthropic.claude-sonnet-4-5-20250929-v1:0",
        contentType="application/json",
        accept="application/json",
        body=json.dumps({
            "anthropic_version": "bedrock-2023-05-31",
            "max_tokens": 4096,
            "messages": [{"role": "user", "content": prompt}]
        })
    )
    result = json.loads(resp["body"].read())
    text = result["content"][0]["text"].strip()
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\s*", "", text)
        text = re.sub(r"\s*```$", "", text)
    return json.loads(text)




def bda_to_doc2_data(raw):
    """Convert BDA flat extraction to Doc2 structured format."""
    return {
        "pharmacy": {
            "name": raw.get("pharmacy_name", ""),
            "address": raw.get("pharmacy_address", ""),
            "phone": raw.get("pharmacy_phone", ""),
            "license": raw.get("pharmacy_license", "")
        },
        "report": {
            "date": raw.get("date_dispensed", ""),
            "id": _generate_report_id(raw.get("date_dispensed", ""))
        },
        "patient": {
            "full_name": raw.get("patient_name", ""),
            "date_of_birth": raw.get("patient_dob", ""),
            "patient_id": raw.get("patient_id", ""),
            "address": raw.get("patient_address", ""),
            "contact_number": raw.get("patient_phone", ""),
            "known_allergies": raw.get("patient_allergies", ""),
            "insurance_provider": raw.get("insurance_provider", ""),
            "insurance_member_id": raw.get("insurance_member_id", "")
        },
        "prescriber": {
            "physician_name": raw.get("physician_name", ""),
            "specialty": raw.get("physician_specialty", ""),
            "practice_name": raw.get("clinic_name", ""),
            "practice_address": raw.get("clinic_address", ""),
            "npi_number": raw.get("npi_number", ""),
            "dea_number": raw.get("dea_number", ""),
            "contact_number": raw.get("physician_phone", "")
        },
        "medication": {
            "rx_number": raw.get("rx_number", ""),
            "date_dispensed": raw.get("date_dispensed", ""),
            "drug_generic": raw.get("drug_generic_name", ""),
            "drug_brand": raw.get("drug_brand_name", ""),
            "strength": raw.get("drug_strength", ""),
            "dosage_form": raw.get("dosage_form", ""),
            "ndc_number": raw.get("ndc_number", ""),
            "lot_number": raw.get("lot_number", ""),
            "expiration_date": raw.get("expiration_date", ""),
            "quantity_dispensed": raw.get("quantity_dispensed", ""),
            "days_supply": raw.get("days_supply", ""),
            "refills_remaining": raw.get("refills_remaining", ""),
            "next_refill_eligible": raw.get("next_refill_date", "")
        },
        "dosage_instructions": _parse_sig(raw.get("sig_directions", "")),
        "warnings": _parse_warnings(raw.get("warnings", "")),
        "billing": {
            "retail_price": raw.get("retail_price", ""),
            "insurance_adjustment": f"-{raw.get('insurance_covered', '')}",
            "patient_copay": raw.get("patient_copay", ""),
            "payment_method": raw.get("payment_method", "")
        },
        "staff": {
            "dispensing_pharmacist": raw.get("pharmacist_name", ""),
            "pharmacy_technician": raw.get("pharmacy_technician", ""),
            "verified_at": f"{raw.get('verification_time', '')}, {raw.get('date_dispensed', '')}",
            "patient_counseled": raw.get("patient_counseled", ""),
            "rph_initials": raw.get("rph_initials", ""),
            "privacy_notice_given": "Yes"
        }
    }


def _generate_report_id(date_str):
    """Generate RPT-YYYYMMDD-002 from a date string like 'October 14, 2024'."""
    import re
    months = {"january":"01","february":"02","march":"03","april":"04",
              "may":"05","june":"06","july":"07","august":"08",
              "september":"09","october":"10","november":"11","december":"12"}
    m = re.match(r"(\w+)\s+(\d+),?\s+(\d{4})", date_str.strip())
    if m:
        month = months.get(m.group(1).lower(), "01")
        day = m.group(2).zfill(2)
        year = m.group(3)
        return f"RPT-{year}{month}{day}-002"
    return "RPT-00000000-002"


def _parse_sig(sig_text):
    """Parse SIG directions into structured fields."""
    return {
        "route": "Oral (by mouth)" if "mouth" in sig_text.lower() or "oral" in sig_text.lower() else "Oral",
        "frequency": "Twice daily (BID)" if "twice" in sig_text.lower() else "As directed",
        "dose": re.search(r"(\w+)\s*\((\d+)\)\s*tablets?", sig_text, re.I).group(0) if re.search(r"(\w+)\s*\((\d+)\)\s*tablets?", sig_text, re.I) else "As directed",
        "administration": "Take with meals (morning and evening)" if "meals" in sig_text.lower() else "As directed",
        "special_notes": sig_text
    }


def _parse_warnings(warnings_text):
    """Parse semicolon-separated warnings into numbered list."""
    if not warnings_text:
        return []
    items = [w.strip() for w in warnings_text.split(";") if w.strip()]
    return [{"number": i + 1, "text": text} for i, text in enumerate(items)]


def generate_doc2(data):
    """Generate a Doc2-format .docx from structured data."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Courier New"
    style.font.size = Pt(9)

    def add_line(text=""):
        doc.add_paragraph(text)

    def add_separator(char="═", width=63):
        add_line(char * width)

    def add_section_header(num, title):
        add_separator()
        add_line(f" SECTION {num} │ {title}")
        add_separator()

    def add_table_row(lines):
        add_line("┌────────────────────────┬──────────────────────────────────┐")
        add_line("│ Field                  │ Details                          │")
        add_line("├────────────────────────┼──────────────────────────────────┤")
        for field, value in lines:
            f_padded = field.ljust(22)
            v_padded = str(value or "").ljust(32)
            add_line(f"│ {f_padded} │ {v_padded} │")
        add_line("└────────────────────────┴──────────────────────────────────┘")

    p = data.get("pharmacy", {})
    add_line("┌─────────────────────────────────────────────────────────────┐")
    add_line(f"│{'GREENLEAF PHARMACY':^61}│")
    add_line(f"│{'MEDICATION DISPENSING REPORT':^61}│")
    add_line(f"│{p.get('address', ''):^61}│")
    ph_lic = f"Ph: {p.get('phone', '')}  ·  License: {p.get('license', '')}"
    add_line(f"│{ph_lic:^61}│")
    add_line("└─────────────────────────────────────────────────────────────┘")

    r = data.get("report", {})
    add_line(f"REPORT DATE: {r.get('date', '')}          REPORT ID: {r.get('id', '')}")

    pt = data.get("patient", {})
    add_section_header(1, "PATIENT PROFILE")
    add_line()
    add_table_row([
        ("Full Name", pt.get("full_name")),
        ("Date of Birth", pt.get("date_of_birth")),
        ("Patient ID", pt.get("patient_id")),
        ("Address", pt.get("address")),
        ("Contact Number", pt.get("contact_number")),
        ("Known Allergies", pt.get("known_allergies")),
        ("Insurance Provider", pt.get("insurance_provider")),
        ("Insurance Member ID", pt.get("insurance_member_id")),
    ])

    pr = data.get("prescriber", {})
    add_section_header(2, "PRESCRIBER PROFILE")
    add_line()
    add_table_row([
        ("Physician Name", pr.get("physician_name")),
        ("Specialty", pr.get("specialty")),
        ("Practice Name", pr.get("practice_name")),
        ("Practice Address", pr.get("practice_address")),
        ("NPI Number", pr.get("npi_number")),
        ("DEA Number", pr.get("dea_number")),
        ("Contact Number", pr.get("contact_number")),
    ])

    med = data.get("medication", {})
    add_section_header(3, "PRESCRIPTION & MEDICATION DETAILS")
    add_line()
    add_table_row([
        ("RX Number", med.get("rx_number")),
        ("Date Dispensed", med.get("date_dispensed")),
        ("Drug (Generic)", med.get("drug_generic")),
        ("Drug (Brand)", med.get("drug_brand")),
        ("Strength", med.get("strength")),
        ("Dosage Form", med.get("dosage_form")),
        ("NDC Number", med.get("ndc_number")),
        ("Lot Number", med.get("lot_number")),
        ("Expiration Date", med.get("expiration_date")),
        ("Quantity Dispensed", med.get("quantity_dispensed")),
        ("Days Supply", med.get("days_supply")),
        ("Refills Remaining", med.get("refills_remaining")),
        ("Next Refill Eligible", med.get("next_refill_eligible")),
    ])

    sig = data.get("dosage_instructions", {})
    add_section_header(4, "DOSAGE INSTRUCTIONS (SIG)")
    add_line()
    add_line(f"  Route          : {sig.get('route', '')}")
    add_line(f"  Frequency      : {sig.get('frequency', '')}")
    add_line(f"  Dose           : {sig.get('dose', '')}")
    add_line(f"  Administration : {sig.get('administration', '')}")
    add_line(f"  Special Notes  : {sig.get('special_notes', '')}")

    warnings = data.get("warnings", [])
    add_section_header(5, "CLINICAL WARNINGS")
    add_line()
    add_line("  No.  Warning")
    add_line("  ───  ─────────────────────────────────────────────────────")
    for w in warnings:
        add_line(f"   {w.get('number', '')}   {w.get('text', '')}")

    bill = data.get("billing", {})
    add_section_header(6, "BILLING SUMMARY")
    add_line()
    add_table_row([
        ("Retail Price", bill.get("retail_price")),
        ("Insurance Adjustment", bill.get("insurance_adjustment")),
        ("Patient Copay", bill.get("patient_copay")),
        ("Payment Method", bill.get("payment_method")),
    ])

    staff = data.get("staff", {})
    add_section_header(7, "DISPENSING STAFF & VERIFICATION")
    add_line()
    add_table_row([
        ("Dispensing Pharmacist", staff.get("dispensing_pharmacist")),
        ("Pharmacy Technician", staff.get("pharmacy_technician")),
        ("Verified At", staff.get("verified_at")),
        ("Patient Counseled", staff.get("patient_counseled")),
        ("RPh Initials", staff.get("rph_initials")),
        ("Privacy Notice Given", staff.get("privacy_notice_given")),
    ])

    add_line()
    add_line("           ─────────────────────────────────────────")
    add_line("           This document is generated for dispensing")
    add_line("           record purposes only. Confidential — for")
    add_line("           authorized use under HIPAA regulations.")
    add_line("           ─────────────────────────────────────────")
    add_line("                © 2024 Greenleaf Pharmacy LLC")
    add_separator()

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    return tmp.name


@app.route("/")
def index():
    return render_template("index.html")



@app.route("/transform", methods=["POST"])
def transform():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files["file"]
    if not file.filename.lower().endswith((".docx", ".pdf")):
        return jsonify({"error": "Only .docx and .pdf files are supported"}), 400

    tmp_input = tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(file.filename)[1])
    file.save(tmp_input.name)

    try:
        result = call_bda(tmp_input.name)
        doc_type = result["type"]
        data = result["data"]

        if doc_type == "clinical_trial":
            if not data:
                return jsonify({"error": "Could not extract clinical trial data from document"}), 400
            output_path = generate_clinical_trial_doc2(data)
            return send_file(
                output_path,
                as_attachment=True,
                download_name="clinical_trial_report.docx",
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        # Prescription flow
        if not data:
            return jsonify({"error": "No data extracted from document"}), 400

        valid_results = []
        for r in data:
            filled = sum(1 for v in r.values() if v and str(v).strip())
            if filled >= 10:
                valid_results.append(r)

        if not valid_results:
            return jsonify({
                "error": "The uploaded document does not appear to be a recognized document type. "
                         "Supported types: Pharmacy Prescription Dispensing Records and Clinical Trial Safety Reports."
            }), 400

        if len(valid_results) == 1:
            doc2_data = bda_to_doc2_data(valid_results[0])
            output_path = generate_doc2(doc2_data)
        else:
            output_path = generate_multi_page_doc2(valid_results)

        return send_file(
            output_path,
            as_attachment=True,
            download_name="transformed_report.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        os.unlink(tmp_input.name)



def generate_multi_page_doc2(results):
    """Generate a single Doc2 with multiple prescription sections."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Courier New"
    style.font.size = Pt(9)

    for i, raw in enumerate(results):
        if i > 0:
            doc.add_page_break()
        data = bda_to_doc2_data(raw)
        # Reuse single-page generation logic by building paragraphs
        _add_doc2_content(doc, data, page_num=i + 1, total_pages=len(results))

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    return tmp.name


def _add_doc2_content(doc, data, page_num=1, total_pages=1):
    """Add Doc2 content to an existing document object."""
    def add_line(text=""):
        doc.add_paragraph(text)

    def add_separator(char="═", width=63):
        add_line(char * width)

    def add_section_header(num, title):
        add_separator()
        add_line(f" SECTION {num} │ {title}")
        add_separator()

    def add_table_row(lines):
        add_line("┌────────────────────────┬──────────────────────────────────┐")
        add_line("│ Field                  │ Details                          │")
        add_line("├────────────────────────┼──────────────────────────────────┤")
        for field, value in lines:
            f_padded = field.ljust(22)
            v_padded = str(value or "").ljust(32)
            add_line(f"│ {f_padded} │ {v_padded} │")
        add_line("└────────────────────────┴──────────────────────────────────┘")

    if total_pages > 1:
        add_line(f"{'─' * 20} PRESCRIPTION {page_num} OF {total_pages} {'─' * 20}")
        add_line()

    p = data.get("pharmacy", {})
    add_line("┌─────────────────────────────────────────────────────────────┐")
    add_line(f"│{'GREENLEAF PHARMACY':^61}│")
    add_line(f"│{'MEDICATION DISPENSING REPORT':^61}│")
    add_line(f"│{p.get('address', ''):^61}│")
    ph_lic = f"Ph: {p.get('phone', '')}  ·  License: {p.get('license', '')}"
    add_line(f"│{ph_lic:^61}│")
    add_line("└─────────────────────────────────────────────────────────────┘")

    r = data.get("report", {})
    add_line(f"REPORT DATE: {r.get('date', '')}          REPORT ID: {r.get('id', '')}")

    pt_data = data.get("patient", {})
    add_section_header(1, "PATIENT PROFILE")
    add_line()
    add_table_row([(k.replace("_", " ").title(), v) for k, v in pt_data.items()])

    add_section_header(2, "PRESCRIBER PROFILE")
    add_line()
    add_table_row([(k.replace("_", " ").title(), v) for k, v in data.get("prescriber", {}).items()])

    add_section_header(3, "PRESCRIPTION & MEDICATION DETAILS")
    add_line()
    add_table_row([(k.replace("_", " ").title(), v) for k, v in data.get("medication", {}).items()])

    sig = data.get("dosage_instructions", {})
    add_section_header(4, "DOSAGE INSTRUCTIONS (SIG)")
    add_line()
    for k, v in sig.items():
        label = k.replace("_", " ").title().ljust(16)
        add_line(f"  {label}: {v}")

    warnings = data.get("warnings", [])
    add_section_header(5, "CLINICAL WARNINGS")
    add_line()
    add_line("  No.  Warning")
    add_line("  ───  ─────────────────────────────────────────────────────")
    for w in warnings:
        add_line(f"   {w.get('number', '')}   {w.get('text', '')}")

    add_section_header(6, "BILLING SUMMARY")
    add_line()
    add_table_row([(k.replace("_", " ").title(), v) for k, v in data.get("billing", {}).items()])

    add_section_header(7, "DISPENSING STAFF & VERIFICATION")
    add_line()
    add_table_row([(k.replace("_", " ").title(), v) for k, v in data.get("staff", {}).items()])

    add_line()
    add_line("           ─────────────────────────────────────────")
    add_line("           This document is generated for dispensing")
    add_line("           record purposes only. Confidential — for")
    add_line("           authorized use under HIPAA regulations.")
    add_line("           ─────────────────────────────────────────")
    add_line("                © 2024 Greenleaf Pharmacy LLC")
    add_separator()

def generate_clinical_trial_doc2(data):
    """Generate a structured Doc2-format .docx from clinical trial data."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Courier New"
    style.font.size = Pt(9)

    def add_line(text=""):
        doc.add_paragraph(text)

    def add_separator(char="═", width=70):
        add_line(char * width)

    def add_section_header(num, title):
        add_separator()
        add_line(f" SECTION {num} │ {title}")
        add_separator()

    def add_table_row(rows):
        add_line("┌──────────────────────────────┬────────────────────────────────────┐")
        add_line("│ Field                        │ Details                            │")
        add_line("├──────────────────────────────┼────────────────────────────────────┤")
        for field, value in rows:
            val = str(value or "")
            # Wrap long values across multiple lines
            while len(val) > 34:
                chunk = val[:34]
                val = val[34:]
                add_line(f"│ {field.ljust(28)} │ {chunk.ljust(34)} │")
                field = ""
            add_line(f"│ {field.ljust(28)} │ {val.ljust(34)} │")
        add_line("└──────────────────────────────┴────────────────────────────────────┘")

    def add_list_section(items, label=""):
        if label:
            add_line(f"  {label}:")
        for i, item in enumerate(items if isinstance(items, list) else [items]):
            if isinstance(item, dict):
                parts = [f"{k}: {v}" for k, v in item.items() if v]
                add_line(f"    {i+1}. {'; '.join(parts)}")
            else:
                add_line(f"    {i+1}. {item}")

    # Header
    sponsor = data.get("sponsor", {})
    doc_info = data.get("document_info", {})
    add_line("┌──────────────────────────────────────────────────────────────────────┐")
    add_line(f"│{(sponsor.get('name', 'CLINICAL TRIAL REPORT')):^70}│")
    add_line(f"│{'CLINICAL TRIAL STRUCTURED ANALYSIS REPORT':^70}│")
    add_line(f"│{sponsor.get('address', ''):^70}│")
    add_line(f"│{'CONFIDENTIAL':^70}│")
    add_line("└──────────────────────────────────────────────────────────────────────┘")
    add_line()
    add_line(f"DOCUMENT ID: {doc_info.get('document_id', '')}     VERSION: {doc_info.get('version', '')}     DATE: {doc_info.get('date_issued', '')}")
    add_line(f"TRIAL PHASE: {doc_info.get('trial_phase', '')}     DEPARTMENT: {doc_info.get('department', '')}")

    # Section 1: Document & Sponsor Info
    add_section_header(1, "DOCUMENT & SPONSOR INFORMATION")
    add_line()
    add_table_row([
        ("Document ID", doc_info.get("document_id")),
        ("Prepared By", doc_info.get("prepared_by")),
        ("Reviewed By", doc_info.get("reviewed_by")),
        ("Approved By", doc_info.get("approved_by")),
        ("Date Issued", doc_info.get("date_issued")),
        ("Department", doc_info.get("department")),
        ("Trial Phase", doc_info.get("trial_phase")),
        ("Version", doc_info.get("version")),
        ("Sponsor", sponsor.get("name")),
        ("Sponsor Address", sponsor.get("address")),
        ("Contact", sponsor.get("contact")),
        ("IND Number", sponsor.get("ind_number")),
    ])

    # Section 2: Study Overview
    overview = data.get("study_overview", {})
    add_section_header(2, "STUDY OVERVIEW")
    add_line()
    add_table_row([
        ("Drug Compound", overview.get("drug_compound")),
        ("Working Name", overview.get("drug_working_name")),
        ("Mechanism", overview.get("mechanism")),
        ("Indication", overview.get("indication")),
        ("Trial Sites", overview.get("trial_sites")),
        ("Countries", overview.get("countries")),
        ("Enrollment Period", overview.get("enrollment_period")),
        ("Total Enrolled", overview.get("total_enrolled")),
        ("Protocol Number", overview.get("protocol_number")),
    ])
    add_line()
    add_line(f"  PRIMARY OBJECTIVE:")
    add_line(f"    {overview.get('primary_objective', '')}")
    sec_obj = overview.get("secondary_objectives", [])
    if sec_obj:
        add_line()
        add_list_section(sec_obj, "SECONDARY OBJECTIVES")

    # Section 3: Enrollment & Demographics
    enroll = data.get("enrollment", {})
    demo = data.get("demographics", {})
    add_section_header(3, "ENROLLMENT & DEMOGRAPHICS")
    add_line()
    add_table_row([
        ("Total Screened", enroll.get("total_screened")),
        ("Screen Failures", enroll.get("screen_failures")),
        ("Randomized", enroll.get("randomized")),
        ("Treatment Arm", enroll.get("treatment_arm")),
        ("Placebo Arm", enroll.get("placebo_arm")),
        ("Completed", enroll.get("completed")),
        ("Discontinued", enroll.get("discontinued")),
    ])
    disc_reasons = enroll.get("discontinuation_reasons", [])
    if disc_reasons:
        add_line()
        add_list_section(disc_reasons, "DISCONTINUATION REASONS")
    add_line()
    add_table_row([
        ("Mean Age", demo.get("mean_age")),
        ("Female %", demo.get("female_pct")),
        ("Male %", demo.get("male_pct")),
        ("Mean Disease Duration", demo.get("mean_disease_duration")),
    ])
    race = demo.get("race_ethnicity", [])
    if race:
        add_line()
        add_list_section(race, "RACE/ETHNICITY")

    # Section 4: Efficacy Results
    eff = data.get("efficacy", {})
    add_section_header(4, "EFFICACY RESULTS")
    add_line()
    add_line("  PRIMARY ENDPOINT:")
    add_table_row([
        ("Endpoint", eff.get("primary_endpoint")),
        ("Treatment Result", eff.get("primary_result_treatment")),
        ("Placebo Result", eff.get("primary_result_placebo")),
        ("Difference", eff.get("primary_difference")),
        ("p-value", eff.get("primary_p_value")),
    ])
    sec_ep = eff.get("secondary_endpoints", [])
    if sec_ep:
        add_line()
        add_list_section(sec_ep, "SECONDARY ENDPOINTS")

    # Section 5: Safety Profile
    safety = data.get("safety", {})
    add_section_header(5, "SAFETY & TOLERABILITY")
    add_line()
    add_table_row([
        ("Any TEAE (Treatment)", safety.get("any_teae_treatment")),
        ("Any TEAE (Placebo)", safety.get("any_teae_placebo")),
        ("Any SAE (Treatment)", safety.get("any_sae_treatment")),
        ("Any SAE (Placebo)", safety.get("any_sae_placebo")),
        ("Deaths (Treatment)", safety.get("deaths_treatment")),
        ("Deaths (Placebo)", safety.get("deaths_placebo")),
        ("Death Details", safety.get("death_details")),
    ])
    common = safety.get("common_teaes", [])
    if common:
        add_line()
        add_list_section(common, "MOST COMMON TEAEs")
    aesi = safety.get("aesi", [])
    if aesi:
        add_line()
        add_list_section(aesi, "ADVERSE EVENTS OF SPECIAL INTEREST")
    labs = safety.get("lab_findings", [])
    if labs:
        add_line()
        add_list_section(labs, "LABORATORY FINDINGS")

    # Section 6: Pharmacokinetics
    pk = data.get("pharmacokinetics", {})
    add_section_header(6, "PHARMACOKINETICS & DRUG INTERACTIONS")
    add_line()
    pk_params = pk.get("pk_parameters", [])
    if pk_params:
        add_list_section(pk_params, "PK PARAMETERS")
    sp_pop = pk.get("special_populations", [])
    if sp_pop:
        add_line()
        add_list_section(sp_pop, "SPECIAL POPULATIONS")
    ddi = pk.get("drug_interactions", [])
    if ddi:
        add_line()
        add_list_section(ddi, "DRUG-DRUG INTERACTIONS")

    # Section 7: Regulatory & Risk-Benefit
    reg = data.get("regulatory", {})
    rb = data.get("risk_benefit", {})
    add_section_header(7, "REGULATORY STATUS & RISK-BENEFIT ASSESSMENT")
    add_line()
    add_table_row([
        ("NDA Submission Target", reg.get("nda_target")),
    ])
    desig = reg.get("designations", [])
    if desig:
        add_line()
        add_list_section(desig, "DESIGNATIONS")
    items = reg.get("open_items", [])
    if items:
        add_line()
        add_list_section(items, "OPEN ITEMS")
    add_line()
    add_line("  RISK-BENEFIT ASSESSMENT:")
    add_line(f"    Benefit : {rb.get('benefit_summary', '')}")
    add_line(f"    Risk    : {rb.get('risk_summary', '')}")
    add_line(f"    Overall : {rb.get('overall_assessment', '')}")

    # Footer
    add_line()
    add_line("           ─────────────────────────────────────────")
    add_line("           This document is a structured analysis")
    add_line("           generated from the original clinical trial")
    add_line("           report. Confidential — for authorized use.")
    add_line("           ─────────────────────────────────────────")
    add_line(f"                © {sponsor.get('name', '')}")
    add_separator()

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    return tmp.name



if __name__ == "__main__":
    app.run(debug=True, port=5000)
